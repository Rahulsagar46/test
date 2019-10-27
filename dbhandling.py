# -*- coding: utf-8 -*-
"""
Created on Sat Oct 26 15:26:16 2019

@author: Rahul
"""
from sqlalchemy.orm import sessionmaker
from bdschema import Application,engine
import core.basicfileops as bfp
from docx import Document


class ManageApplications:
    
    def __init__(self,company,jobid,role,location):
        global session  
        global newapplication
        
        self.company = company
        self.jobid = jobid
        self.role = role
        self.location = location
        
        if(self.jobid == ""):
            self.jobid = "None"
        if(self.location == ""):
            self.location = "None"
            
        Session = sessionmaker(bind=engine)
        session =Session()
        newapplication = Application()

    def addEntry(self):
        newapplication.company = self.company
        newapplication.jobid = self.jobid
        newapplication.role = self.role
        newapplication.location = self.location
        
        session.add(newapplication)
        session.commit()
        session.close

        folderStructure = self.company+'/'+self.location+'-'+self.role+'-'+self.jobid
        bfp.createFoldersInDirectory(folderStructure)
        
    def checkSimilarEntries(self):
        check_results = session.query(Application).filter(Application.company == self.company,Application.jobid == self.jobid).count()
        if(check_results > 0):
            print('Similar Entries Exist in Database')
            returnTxt = True
        else:
            returnTxt = False
        
        return returnTxt
    
    def addCoverLetter2Folders(self):
        sourceLocation = 'templates/IoT Consultant.docx'
        defaultLocation = 'C:/Users/Rahul/Desktop/MyStaY aT emdEn/Full time/Applied companies/Phase-2'
        destinationLocation = defaultLocation+'/'+self.company+'/'+self.location+'-'+self.role+'-'+self.jobid
        bfp.copyFile2Location(sourceLocation,destinationLocation)
                                
        document = Document(destinationLocation+'/IoT Consultant.docx')
    
        for paragraph in document.paragraphs:                     
            if '?????????' in paragraph.text or '$$$$$$$$$' in paragraph.text :
#                print(paragraph.text)
                text1 = str(paragraph.text)
                split_txt =text1.split()
                
                final_para = ""
                for txt in split_txt:
                    if(txt == "$$$$$$$$$." or txt == "$$$$$$$$$" or txt == "$$$$$$$$$,"):
                        txt = self.role
                        print('Yes role')
                        
                    elif(txt == '?????????' or txt == '?????????.' or txt == '?????????,'):
                        txt = self.company
                        print('Yes company')
        
                    final_para = final_para+""+txt
                document.add_paragraph(final_para)
                document.save(destinationLocation+'/IoT Consultant.docx')
                        
            else:
                document.add_paragraph(paragraph.text)
                document.save(destinationLocation+'/'+self.company+'-'+self.role+'.docx')
                        
#                result = text1.find('$$$$$$$$$')
#                print(result,text1[52],text1[53])
#                paragraph.text = 'new text containing ocean'
        